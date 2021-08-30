import re
import os
import glob
import docx

target = "*.txt"

def header_footer(document, testname):
    section = document.sections[0]
    header = section.header
    heading = header.paragraphs[0]
    heading.text = testname + "\t\tName:_______________________"

    footer = section.footer
    footing = footer.paragraphs[0]
    run = footing.add_run()
    run.font.size = docx.shared.Pt(8)
    footing.text = "\tCopyright ©2020 Cyber Innovation Center\n\tAll Rights Reserved. Not for Distribution."
    
def question(document, ques_number, ques_text, answers, correct):
    text = str(ques_number)+ ". " + ques_text
    para = document.add_paragraph(text)
    
    tab_stops = para.paragraph_format.tab_stops
    tab_stops.add_tab_stop(docx.shared.Inches(0.5))
    answer_counter = 65
    for answer in answers:
        answer_text = "\n\t"+chr(answer_counter)+". "+answer
        para.add_run(answer_text)
        answer_counter += 1
        
    #correct:
    para.add_run("\n\nAnswer: ________")

for filename_orig in glob.glob(target):
    file_orig = open(filename_orig, "r+", encoding="utf8", errors="ignore")
    filename_new = re.sub(r"(\d+\.\d+\.\d+|Appendix [A-Z])",r"\1 - quiz", filename_orig)
    filename_new = re.sub(r".txt",r".docx", filename_new)
    testname = re.sub(r"(\d+\.\d+\.\d+) - ","", filename_orig)
    testname = re.sub(" - Key", "", testname)
    testname = re.sub(".txt", "", testname)
    #print(filename_new)
    
    #DOCUMENT SETUP
    document = docx.Document()
    font = document.styles['Normal'].font
    font.name = "Arial"
    font.size = docx.shared.Pt(10)
    
    header_footer(document, testname)
    
    question_text = ""
    question_number = 1
    answers = []
    correct = []
    #line_num = 0

#-------------------------------------------------------------------------------
    for line in file_orig:
        #line_num+=1
        #line = unicode(line, errors='ignore')
        line = line.strip()
        #print(line)
        try:
            line = line.encode('ascii','ignore').decode()
        except:
            pass
        #print(line_num, line)
        if re.match("^True", line):
            answers.append("True")
            continue
        if re.match("^False", line):
            answers.append("False")
            continue
        if re.match("^\*True", line):
            answers.append("True")
            correct.append("A")
            continue
        if re.match("^\*False", line):
            answers.append("False")
            correct.append("B")
            continue
        if re.match("^\*[A-Za-z]{1}", line):
            option = re.split("\)|\.|\ ", line, 1)
            #print(option)
            answers.append(option[1].strip())
            correct.append( re.sub("\*","", option[0]) )
            continue
        if re.match("^[A-Za-z]{1}", line):
            option = re.split("\)|\.|\ ", line, 1)
            #print(option)
            answers.append(option[1].strip())
            continue
        #QUESTION TEXT
        if re.match("^[0-9]", line):
            text = line.split(". ", 1)
            question_text = str(text[1])
            #print(text[1])
            continue
        if re.match("", line):
            #print("#"+ str(question_number) +": ", question_text,"\n", answers,"\n", correct)
            #XML output
            if (len(question_text)>0) and (len(answers)>0) and (len(correct)>0):
                question(document, str(question_number), question_text, answers, correct)
                question_number += 1
            question_text = ""
            answers = []
            correct = []
            #print("---")
    
    document.save(filename_new)
    print(filename_new , "created")
    
    file_orig.close()
    print("\n")